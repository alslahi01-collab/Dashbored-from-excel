import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
import io
import zipfile
from rapidfuzz import process, utils

# إعدادات الصفحة
st.set_page_config(page_title="محلل البيانات الذكي", layout="wide")

# دالة توحيد المسميات (Logic Requested)
def unify_names(series):
    counts = series.value_counts()
    unique_names = counts.index.tolist()
    mapping = {}
    
    for name in unique_names:
        if name in mapping: continue
        # البحث عن الأسماء المشابهة بنسبة 80%
        matches = process.extract(name, unique_names, score_cutoff=80)
        group = [m[0] for m in matches]
        
        # اختيار الاسم الأكثر تكراراً
        best_match = counts[group].idxmax()
        # إذا تساوى التكرار، اختيار الأقصر طولاً
        max_freq = counts[group].max()
        best_matches = [n for n in group if counts[n] == max_freq]
        if len(best_matches) > 1:
            best_match = min(best_matches, key=len)
            
        for m in group:
            mapping[m] = best_match
            
    return series.map(mapping)

# عنوان الموقع
st.title("📊 محلل البيانات العربي الذكي")
st.write("قم برفع ملف الإكسل، وسنقوم بالباقي!")

uploaded_file = st.file_uploader("اختر ملف إكسل", type=["xlsx"])

if uploaded_file:
    xl = pd.ExcelFile(uploaded_file)
    sheet_name = st.selectbox("اختر ورقة العمل:", xl.sheet_names)
    df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
    
    st.write("### استعراض البيانات")
    st.dataframe(df.head())

    # اختيار الأعمدة
    all_columns = df.columns.tolist()
    selected_cols = st.multiselect("اختر الأعمدة للتحليل (أو اتركها فارغة لتحليل الكل):", all_columns)
    
    if not selected_cols:
        selected_cols = all_columns

    if st.button("بدء المعالجة والتحليل"):
        with st.spinner('جاري تنظيف وتوحيد البيانات...'):
            df_processed = df[selected_cols].copy()
            
            # عملية التوحيد التلقائي للأعمدة النصية
            for col in df_processed.select_dtypes(include=['object']):
                df_processed[col] = unify_names(df_processed[col])

            st.success("تم توحيد البيانات وتنظيفها!")
            
            # --- قسم الداشبورد ---
            st.header("📈 لوحة المعلومات (Dashboard)")
            cols = st.columns(2)
            charts_images = [] # لتخزين الصور لملف الوورد

            for i, col_name in enumerate(selected_cols):
                target_col = cols[i % 2]
                if df_processed[col_name].dtype in ['int64', 'float64']:
                    # رسم بياني عددي
                    fig = px.histogram(df_processed, x=col_name, title=f"توزيع {col_name}", color_discrete_sequence=['#636EFA'])
                    target_col.plotly_chart(fig)
                else:
                    # رسم بياني فئوي
                    top_values = df_processed[col_name].value_counts().head(10)
                    fig = px.bar(x=top_values.index, y=top_values.values, title=f"أعلى القيم في {col_name}")
                    target_col.plotly_chart(fig)
                
                # حفظ الصورة للتقرير (بشكل مؤقت)
                img_bytes = fig.to_image(format="png")
                charts_images.append((col_name, img_bytes))

            # --- إنشاء التعليقات والتحليل (Textual Analysis) ---
            st.header("📝 التفسير والتحليل")
            analysis_text = []
            for col_name in selected_cols:
                if df_processed[col_name].dtype in ['int64', 'float64']:
                    mean_val = df_processed[col_name].mean()
                    comment = f"العمود '{col_name}': يبلغ متوسط القيم حوالي {mean_val:.2f}. "
                    if df_processed[col_name].skew() > 0:
                        comment += "البيانات تنحو نحو القيم المرتفعة."
                    else:
                        comment += "البيانات موزعة بشكل متوازن نسبيًا."
                else:
                    top_val = df_processed[col_name].mode()[0]
                    comment = f"العمود '{col_name}': القيمة الأكثر تكراراً هي '{top_val}'. "
                
                st.write(f"- {comment}")
                analysis_text.append(comment)

            # --- إنشاء ملفات المخرجات ---
            # 1. ملف إكسل المعالج
            excel_buffer = io.BytesIO()
            df_processed.to_excel(excel_buffer, index=False)
            
            # 2. ملف الوورد
            doc = Document()
            doc.add_heading('تقرير تحليل البيانات التلقائي', 0)
            for text, (col_name, img) in zip(analysis_text, charts_images):
                doc.add_heading(f'تحليل عمود: {col_name}', level=1)
                doc.add_paragraph(text)
                image_stream = io.BytesIO(img)
                doc.add_picture(image_stream, width=Inches(5))
            
            word_buffer = io.BytesIO()
            doc.save(word_buffer)

            # 3. إنشاء ملف ZIP
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.writestr("Cleaned_Data.xlsx", excel_buffer.getvalue())
                zip_file.writestr("Analysis_Report.docx", word_buffer.getvalue())

            st.download_button(
                label="📥 تحميل الملفات (ZIP)",
                data=zip_buffer.getvalue(),
                file_name="Data_Analysis_Package.zip",
                mime="application/zip"
            )